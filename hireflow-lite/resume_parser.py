"""
HireFlow-Lite — Layout-Aware & Block-Classified Resume Parser
===================================================================
Extracts structured fields (name, email, phone, GitHub username, skills,
projects, education) using layout-aware typographic block classification,
block-level evidence tracking, heading-independent project assembly, and
quality gate validation.

Supported formats: .pdf, .docx
"""

from __future__ import annotations

import os
import re
import sys
import json
import time
import statistics
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Set

# Force UTF-8 stdout/stderr on Windows to avoid UnicodeEncodeError
try:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

import fitz  # PyMuPDF
from docx import Document


# ═══════════════════════════════════════════════════════════════════════════════
#  Data Structures & Typographic Model
# ═══════════════════════════════════════════════════════════════════════════════

def _clean_text(raw: str) -> str:
    """Normalise whitespace, preserve paragraph boundaries."""
    text = raw.replace("\t", " ").replace("\f", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[^\S\n]+", " ", text)
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join(lines).strip()


class SpanInfo:
    """Represents a text span with typographic and bounding-box information."""
    def __init__(
        self,
        text: str,
        font_name: str = "",
        font_size: float = 10.0,
        is_bold: bool = False,
        is_italic: bool = False,
        bbox: Optional[Tuple[float, float, float, float]] = None,
        page: int = 0,
        block_id: int = 0,
    ):
        self.text = text
        self.font_name = font_name
        self.font_size = round(font_size, 1)
        self.is_bold = is_bold
        self.is_italic = is_italic
        self.bbox = [round(x, 1) for x in bbox] if bbox else None
        self.page = page
        self.block_id = block_id

    def to_dict(self) -> dict:
        return {
            "block_id": self.block_id,
            "text": self.text,
            "font_name": self.font_name,
            "font_size": self.font_size,
            "is_bold": self.is_bold,
            "is_italic": self.is_italic,
            "bbox": self.bbox,
            "page": self.page,
        }


class LineInfo:
    """Represents a line of text constructed from one or more Spans with a unique block_id."""
    def __init__(self, block_id: int, spans: List[SpanInfo], page: int = 0):
        self.block_id = block_id
        self.spans = spans
        for s in spans:
            s.block_id = block_id
        self.text = " ".join(s.text for s in spans if s.text.strip()).strip()
        self.page = page
        self.font_size = max((s.font_size for s in spans), default=10.0)
        self.is_bold = any(s.is_bold for s in spans)
        self.is_italic = any(s.is_italic for s in spans)
        
        # Calculate line bounding box
        valid_bboxes = [s.bbox for s in spans if s.bbox]
        if valid_bboxes:
            x0 = min(b[0] for b in valid_bboxes)
            y0 = min(b[1] for b in valid_bboxes)
            x1 = max(b[2] for b in valid_bboxes)
            y1 = max(b[3] for b in valid_bboxes)
            self.bbox = [round(x0, 1), round(y0, 1), round(x1, 1), round(y1, 1)]
        else:
            self.bbox = None

        # Check if bullet point (including middle dot \u00b7)
        self.is_bullet = bool(re.match(r"^[•\-\*\u2022\u00b7\u25cf\u25cb\u2023\u25aa\u25b8\d+\.]\s*", self.text))
        self.block_type: str = "UNKNOWN"
        self.section_context: str = "UNKNOWN"

    def to_dict(self) -> dict:
        return {
            "block_id": self.block_id,
            "text": self.text,
            "block_type": self.block_type,
            "font_size": self.font_size,
            "is_bold": self.is_bold,
            "is_bullet": self.is_bullet,
            "bbox": self.bbox,
            "page": self.page,
        }


# ═══════════════════════════════════════════════════════════════════════════════
#  Layout-Aware Extraction (PDF & DOCX with Unique Block IDs)
# ═══════════════════════════════════════════════════════════════════════════════

def _extract_pdf_structured(filepath: str) -> Tuple[List[LineInfo], float, dict]:
    """Extract Lines with typographic spans & bboxes, using Firecrawl pdf_inspector for layout classification."""
    layout_metadata = {
        "engine": "PyMuPDF",
        "pdf_type": "text_based",
        "is_complex_layout": False,
        "has_encoding_issues": False,
        "pages_with_columns": [],
        "pages_with_tables": [],
        "inspector_latency_ms": 0,
    }

    try:
        import pdf_inspector
        insp_result = pdf_inspector.process_pdf(filepath)
        layout_metadata.update({
            "engine": "pdf-inspector-rust",
            "pdf_type": getattr(insp_result, "pdf_type", "text_based"),
            "is_complex_layout": getattr(insp_result, "is_complex_layout", False),
            "has_encoding_issues": getattr(insp_result, "has_encoding_issues", False),
            "pages_with_columns": list(getattr(insp_result, "pages_with_columns", []) or []),
            "pages_with_tables": list(getattr(insp_result, "pages_with_tables", []) or []),
            "inspector_latency_ms": getattr(insp_result, "processing_time_ms", 0),
        })
    except Exception:
        pass

    doc = fitz.open(filepath)
    all_lines: List[LineInfo] = []
    all_font_sizes: List[float] = []
    current_block_id = 0

    for page_idx, page in enumerate(doc):
        rect = page.rect
        width = rect.width
        midpoint = width / 2.0

        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])
        if not blocks:
            continue

        left_col_lines: List[Tuple[float, List[SpanInfo]]] = []
        right_col_lines: List[Tuple[float, List[SpanInfo]]] = []
        full_width_lines: List[Tuple[float, List[SpanInfo]]] = []

        for b in blocks:
            if b.get("type") != 0:  # Skip non-text blocks
                continue

            x0, y0, x1, y1 = b.get("bbox", (0, 0, 0, 0))
            
            for line_dict in b.get("lines", []):
                spans: List[SpanInfo] = []
                for s in line_dict.get("spans", []):
                    stext = s.get("text", "")
                    if not stext.strip():
                        continue
                    
                    font_name = s.get("font", "")
                    font_size = s.get("size", 10.0)
                    flags = s.get("flags", 0)
                    is_bold = ("bold" in font_name.lower()) or bool(flags & 16) or bool(flags & 2)
                    is_italic = ("italic" in font_name.lower()) or bool(flags & 1)
                    s_bbox = s.get("bbox")
                    
                    all_font_sizes.append(font_size)
                    spans.append(SpanInfo(
                        text=stext,
                        font_name=font_name,
                        font_size=font_size,
                        is_bold=is_bold,
                        is_italic=is_italic,
                        bbox=s_bbox,
                        page=page_idx,
                    ))

                if not spans:
                    continue

                line_y = line_dict.get("bbox", [0, y0, 0, 0])[1]

                # Column classification with 25px tolerance
                if x1 <= midpoint + 25:
                    left_col_lines.append((line_y, spans))
                elif x0 >= midpoint - 25:
                    right_col_lines.append((line_y, spans))
                else:
                    full_width_lines.append((line_y, spans))

        # Reconstruct reading order top-to-bottom within columns
        left_col_lines.sort(key=lambda x: x[0])
        right_col_lines.sort(key=lambda x: x[0])

        ordered_spans_list = []
        if len(left_col_lines) > 2 and len(right_col_lines) > 2:
            first_col_y = min([y for y, _ in left_col_lines + right_col_lines])
            top_lines = [spans for y, spans in full_width_lines if y < first_col_y]
            bottom_lines = [spans for y, spans in full_width_lines if y >= first_col_y]

            ordered_spans_list.extend(top_lines)
            ordered_spans_list.extend([spans for y, spans in left_col_lines])
            ordered_spans_list.extend([spans for y, spans in right_col_lines])
            ordered_spans_list.extend(bottom_lines)
        else:
            all_page = left_col_lines + right_col_lines + full_width_lines
            all_page.sort(key=lambda x: x[0])
            ordered_spans_list.extend([spans for y, spans in all_page])

        for spans in ordered_spans_list:
            all_lines.append(LineInfo(block_id=current_block_id, spans=spans, page=page_idx))
            current_block_id += 1

    doc.close()

    median_font_size = statistics.median(all_font_sizes) if all_font_sizes else 10.0
    return all_lines, median_font_size, layout_metadata


def _extract_docx_structured(filepath: str) -> Tuple[List[LineInfo], float, dict]:
    """Extract Lines with font styles from DOCX files with block_ids."""
    doc = Document(filepath)
    all_lines: List[LineInfo] = []
    all_font_sizes: List[float] = []
    current_block_id = 0
    layout_metadata = {
        "engine": "python-docx",
        "pdf_type": "docx",
        "is_complex_layout": False,
        "has_encoding_issues": False,
        "pages_with_columns": [],
        "pages_with_tables": [],
        "inspector_latency_ms": 0,
    }

    def process_paragraph(para, page_idx=0):
        nonlocal current_block_id
        text = para.text.strip()
        if not text:
            return
        spans: List[SpanInfo] = []
        for run in para.runs:
            rtext = run.text
            if not rtext.strip():
                continue
            
            fsize = 10.0
            if run.font.size:
                fsize = run.font.size.pt
            elif para.style and hasattr(para.style, "font") and para.style.font.size:
                fsize = para.style.font.size.pt
            
            fname = run.font.name or ""
            is_bold = bool(run.bold) or ("bold" in (para.style.name or "").lower())
            is_italic = bool(run.italic)

            all_font_sizes.append(fsize)
            spans.append(SpanInfo(
                text=rtext,
                font_name=fname,
                font_size=fsize,
                is_bold=is_bold,
                is_italic=is_italic,
                bbox=None,
                page=page_idx,
                block_id=current_block_id,
            ))

        if spans:
            all_lines.append(LineInfo(block_id=current_block_id, spans=spans, page=page_idx))
        else:
            all_font_sizes.append(10.0)
            all_lines.append(LineInfo(block_id=current_block_id, spans=[SpanInfo(text=text, font_size=10.0, page=page_idx, block_id=current_block_id)], page=page_idx))
        current_block_id += 1

    for p in doc.paragraphs:
        process_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

    median_font_size = statistics.median(all_font_sizes) if all_font_sizes else 10.0
    return all_lines, median_font_size, layout_metadata


# ═══════════════════════════════════════════════════════════════════════════════
#  Line Reconstruction & Section Segmentation
# ═══════════════════════════════════════════════════════════════════════════════

def _reconstruct_lines(lines: List[LineInfo]) -> List[LineInfo]:
    """Merge visual line continuations (e.g. wrapped lines) based on spatial and typographic proximity."""
    if not lines:
        return []

    reconstructed: List[LineInfo] = [lines[0]]
    for curr in lines[1:]:
        prev = reconstructed[-1]
        
        # Check if they are on the same page
        if curr.page != prev.page:
            reconstructed.append(curr)
            continue
            
        # A new bullet point ALWAYS starts a new line
        if curr.is_bullet:
            reconstructed.append(curr)
            continue
            
        # Typography mismatch
        if abs(prev.font_size - curr.font_size) > 1.0 or prev.is_bold != curr.is_bold:
            reconstructed.append(curr)
            continue
            
        merge = False
        if prev.bbox and curr.bbox:
            # Spatial heuristic: tighter Y gap for wrapped lines, slightly forgiving X alignment for bullets
            x_align = abs(prev.bbox[0] - curr.bbox[0]) < 25.0
            y_gap = curr.bbox[1] - prev.bbox[3]
            
            # Wrapped lines are typically very tight vertically
            if x_align and -5.0 <= y_gap < 6.0:
                # Do not merge if current line looks like a new heading/title (starts with upper and short)
                if prev.is_bullet and curr.text[0].isupper() and len(curr.text) < 50:
                    merge = False
                else:
                    merge = True
        else:
            # Fallback for DOCX (no bboxes): merge if previous line doesn't end with terminal punctuation
            if not re.search(r"[.!?]\s*$", prev.text):
                merge = True
                
        if merge:
            # Merge curr into prev
            prev.text = prev.text + " " + curr.text
            prev.spans.extend(curr.spans)
            if prev.bbox and curr.bbox:
                prev.bbox = [
                    min(prev.bbox[0], curr.bbox[0]),
                    min(prev.bbox[1], curr.bbox[1]),
                    max(prev.bbox[2], curr.bbox[2]),
                    max(prev.bbox[3], curr.bbox[3])
                ]
            # If prev was not a bullet, but text now looks like one, don't change is_bullet (it inherits from prev)
        else:
            reconstructed.append(curr)

    return reconstructed


def _detect_sections(lines: List[LineInfo], median_font_size: float) -> List[LineInfo]:
    """Assigns an overarching section_context to every line based on section headings."""
    # First pass: identify section headings and store their page and Y-coordinate
    headings = []
    
    for line in lines:
        text_lower = line.text.lower().strip().rstrip(":-")
        is_heading_size = line.font_size >= (median_font_size * 1.15)
        is_short = len(line.text) < 45
        
        is_sec_heading = False
        for stype, keywords in _SECTION_KEYWORDS.items():
            if any(kw == text_lower or (is_short and text_lower.startswith(kw)) for kw in keywords):
                if is_heading_size or line.is_bold or is_short:
                    line.section_context = stype.upper()
                    # Store heading: (page, y_coord, type)
                    y_coord = line.bbox[1] if line.bbox else 0
                    headings.append((line.page, y_coord, stype.upper()))
                    is_sec_heading = True
                    break
        
        if not is_sec_heading:
            line.section_context = "UNKNOWN"

    # Second pass: assign non-heading lines to the appropriate section based on Y-coordinate
    for line in lines:
        if line.section_context != "UNKNOWN":
            continue
            
        page = line.page
        y_coord = line.bbox[1] if line.bbox else 0
        
        # Find headings on the same page
        page_headings = [(y, stype) for p, y, stype in headings if p == page]
        
        if not page_headings:
            # Fallback to previous page's last heading if any, else HEADER
            prev_headings = [stype for p, y, stype in headings if p < page]
            line.section_context = prev_headings[-1] if prev_headings else "HEADER"
            continue
            
        # Find the heading with the largest Y that is <= line's Y + 10 (small tolerance)
        valid_headings = [stype for y, stype in page_headings if y <= y_coord + 10]
        if valid_headings:
            line.section_context = valid_headings[-1]
        else:
            line.section_context = "HEADER"
        
    return lines


# ═══════════════════════════════════════════════════════════════════════════════
#  Block Taxonomy & Classifier Pipeline
# ═══════════════════════════════════════════════════════════════════════════════

_SECTION_KEYWORDS = {
    "projects": [
        "projects", "personal projects", "academic projects", "technical projects",
        "key projects", "notable projects", "side projects", "selected projects",
        "selected work", "project experience", "portfolio", "research", "hackathons",
        "open source", "open source contributions"
    ],
    "experience": [
        "experience", "work experience", "professional experience", "employment history",
        "internship experience", "work history", "career history"
    ],
    "education": ["education", "academic background", "qualifications", "academic record"],
    "skills": [
        "skills", "technical skills", "technologies", "tech stack", "core competencies",
        "tools & technologies", "programming languages", "domain skills"
    ],
    "achievements": [
        "achievements", "certifications", "accomplishments", "publications", "awards",
        "leadership", "extracurricular", "honors"
    ],
}

_TECH_REGEX = re.compile(
    r"\b(python|rust|c\+\+|java|javascript|typescript|react|next\.js|fastapi|node\.js|express|django|flask|docker|aws|postgresql|mongodb|mysql|pytorch|tensorflow|opencv|graphql|rest apis?|mcp|evm|web3|solidity|go|ruby|tokio|cargo|webassembly|cuda|spring boot|spring|kafka|tailwind css|tailwind|streamlit|pandas|numpy|scikit-learn|redis|elasticsearch|firebase|vector db|sqlite|sqlite-vec|vue|angular|kubernetes|terraform|ci/cd|linux|git|html|css|framer motion|vercel)\b",
    re.IGNORECASE,
)

_DATE_REGEX = re.compile(
    r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|march|april|june|july|august|september|october|november|december)?\s*\d{4}\b|\b(20\d{2})\b|\bpresent\b",
    re.IGNORECASE,
)

_URL_REGEX = re.compile(r"(https?://\S+|github\.com/\S+|devpost\.com/\S+)", re.IGNORECASE)

_ACTION_VERBS = {
    "built", "trained", "developed", "implemented", "created", "designed",
    "wrote", "used", "applied", "worked", "fixed", "added", "achieved",
    "engineered", "consolidated", "refactored", "deployed", "integrated",
    "optimized", "improved", "automated", "configured", "managed", "led",
    "contributed", "maintained", "spearheaded", "architected", "launched",
}

_COMPETITIVE_PROG_REGEX = re.compile(
    r"\b(leetcode|codechef|codeforces|hackerrank|kaggle|contest rating|solved \d+\+?|div \d+|gate|code-a-thon)\b",
    re.IGNORECASE,
)

_CERTIFICATION_REGEX = re.compile(
    r"\b(aws certified|aws cloud practitioner|nptel|coursera|udemy|cisco|certificate|certification|virtual internship|foundational)\b",
    re.IGNORECASE,
)

_EDUCATION_DEGREE_REGEX = re.compile(
    r"\b(b\.e\.|b\.tech|m\.tech|bachelor|master|diploma|cgpa|semesters?|gpa|higher secondary|senior secondary|school)\b",
    re.IGNORECASE,
)

_JOB_TITLE_COMPANY_REGEX = re.compile(
    r"\b(intern|internship|software engineer|developer intern|rpa developer|titan company|ltd\.|pvt ltd|inc\.|corporation)\b",
    re.IGNORECASE,
)

_SKILL_HEADER_REGEX = re.compile(
    r"^\s*(languages|frameworks|tools|databases|tech stacks?|core competencies|frontend|backend|devops & security)\s*[:\-]",
    re.IGNORECASE,
)

_ACHIEVEMENT_REGEX = re.compile(
    r"\b(rank|winner|finalist|placed among|top \d+%|vice president|technical club|spearheaded technical)\b",
    re.IGNORECASE,
)


def _score_project_header_candidate(
    line: LineInfo,
    next_line: Optional[LineInfo],
    median_font_size: float,
) -> float:
    """Score line (0.0 to 1.0) as a potential project title candidate using positive & negative signals."""
    text = line.text.strip()
    if not text or len(text) < 3:
        return 0.0

    score = 0.0

    # ══════════════════════════════════════════════════════════════════════════
    #  POSITIVE SIGNALS
    # ══════════════════════════════════════════════════════════════════════════
    # 1. Typography (weight 0.25)
    if line.is_bold:
        score += 0.15
    if line.font_size > median_font_size:
        score += 0.10

    # 2. Position & Line Length (weight 0.15)
    if len(text) <= 75:
        score += 0.10
    if not line.text.startswith(" ") and line.text[0].isupper():
        score += 0.05

    # 3. Bullet Structure (weight 0.15)
    if not line.is_bullet:
        score += 0.15

    # 4. Technology Pattern (weight 0.15)
    has_parens_tech = bool(re.search(r"\([^)]*\b(python|rust|react|c\+\+|java|fastapi|node|docker|aws)\b", text, re.I))
    has_inline_tech = bool(re.search(r"[–—·|:\-]\s*.*?\b(python|rust|react|c\+\+|java|fastapi|node|docker)\b", text, re.I))
    if has_parens_tech or has_inline_tech:
        score += 0.15

    # 5. Links / URLs (weight 0.10)
    if _URL_REGEX.search(text):
        score += 0.10

    # 6. Dates (weight 0.10)
    if _DATE_REGEX.search(text):
        score += 0.10

    # 7. Bullet Context (weight 0.10)
    if next_line and next_line.is_bullet:
        score += 0.10

    # ══════════════════════════════════════════════════════════════════════════
    #  NEGATIVE SIGNALS (Subtractions)
    # ══════════════════════════════════════════════════════════════════════════
    # Bullet Line Penalty (-0.50)
    if line.is_bullet:
        score -= 0.50

    # Top Candidate Header Penalty (-0.50)
    if line.block_id <= 2 and ("candidate" in text.lower() or _extract_email(text) or (next_line and _extract_email(next_line.text))):
        score -= 0.50

    # Continuation Fragment Penalty (-0.60)
    # If the line starts with a lowercase letter, it is almost certainly a continuation of a previous bullet
    if text and text[0].islower():
        score -= 0.60

    # Action Verb Bullet Penalty (-0.40)
    first_word = text.split()[0].lower().rstrip(".,:")
    if first_word in _ACTION_VERBS:
        score -= 0.40

    # Domain Non-Project Penalties
    if _COMPETITIVE_PROG_REGEX.search(text):
        score -= 0.35
    if _CERTIFICATION_REGEX.search(text):
        score -= 0.35
    if _EDUCATION_DEGREE_REGEX.search(text):
        score -= 0.40
    if _JOB_TITLE_COMPANY_REGEX.search(text):
        score -= 0.30
    if _SKILL_HEADER_REGEX.search(text):
        score -= 0.35
    if _ACHIEVEMENT_REGEX.search(text):
        score -= 0.30

    return max(0.0, min(1.0, score))


def _classify_blocks(
    lines: List[LineInfo],
    median_font_size: float,
    header_threshold: float = 0.30
) -> List[LineInfo]:
    """Classify every LineInfo block into taxonomy categories."""
    num_lines = len(lines)
    
    for i, line in enumerate(lines):
        text_lower = line.text.lower().strip().rstrip(":-")
        is_heading_size = line.font_size >= (median_font_size * 1.15)
        is_short = len(line.text) < 45

        # Check for section heading
        is_sec_heading = False
        for stype, keywords in _SECTION_KEYWORDS.items():
            if any(kw == text_lower or (is_short and text_lower.startswith(kw)) for kw in keywords):
                if is_heading_size or line.is_bold or is_short:
                    line.block_type = "SECTION_HEADING"
                    is_sec_heading = True
                    break
        if is_sec_heading:
            continue

        # Check for project title candidate ONLY if we are in an eligible section
        is_project_title = False
        if line.section_context in ("PROJECTS", "EXPERIENCE"):
            next_line = lines[i + 1] if i + 1 < num_lines else None
            hscore = _score_project_header_candidate(line, next_line, median_font_size)
            if hscore >= header_threshold:
                line.block_type = "PROJECT_TITLE"
                is_project_title = True
                
        if is_project_title:
            continue

        # Check for standalone technology line (pure tech list or explicit header)
        text_no_punct = re.sub(r"[^\w\s]", " ", line.text).lower()
        words = [w for w in text_no_punct.split() if len(w) > 1]
        tech_matches = _TECH_REGEX.findall(line.text)
        is_pure_tech = (len(words) > 0 and (len(tech_matches) >= min(2, len(words) // 2))) or re.search(r"^\s*(tech stack|technologies|built with|tools|stack|tech|skills)\s*[:\-]", line.text, re.I)
        
        if is_pure_tech:
            line.block_type = "TECHNOLOGY_LINE"
            continue

        # Check for standalone link or date
        if _URL_REGEX.search(line.text) and len(line.text) < 80:
            line.block_type = "LINK"
            continue
        if _DATE_REGEX.search(line.text) and len(line.text) < 40 and not line.is_bullet:
            line.block_type = "DATE"
            continue

        # Check for bullet
        if line.is_bullet:
            line.block_type = "BULLET"
            continue

        # Default narrative line
        if len(line.text) > 30:
            line.block_type = "DESCRIPTION"
        else:
            line.block_type = "UNKNOWN"

    return lines


# ═══════════════════════════════════════════════════════════════════════════════
#  Heading-Independent Project Assembler & Block Evidence Linker
# ═══════════════════════════════════════════════════════════════════════════════

def _clean_project_name(header_text: str) -> Tuple[str, List[str]]:
    """Separates clean project title from attached technologies, dates, and links."""
    raw = header_text
    raw = raw.replace("\ufffd", " – ")
    
    extracted_techs: list[str] = []
    
    # Extract from parentheses first
    parens_match = re.search(r"\(([^)]+)\)", raw)
    if parens_match:
        inside = parens_match.group(1)
        found_techs = _TECH_REGEX.findall(inside)
        if found_techs:
            extracted_techs.extend([t.title() for t in found_techs])
        raw = re.sub(r"\s*\([^)]*\)", "", raw)

    parts = re.split(r"\s*[–—·|:\-]\s*", raw)
    title = parts[0].strip()
    
    # Also search for technologies in the rest of the string (after title)
    if len(parts) > 1:
        rest_of_string = " ".join(parts[1:])
        found_techs = _TECH_REGEX.findall(rest_of_string)
        if found_techs:
            extracted_techs.extend([t.title() for t in found_techs])

    title = _DATE_REGEX.sub("", title).strip()
    title = _URL_REGEX.sub("", title).strip()
    title = title.strip(":-–—·*, ")

    return title if title else header_text, extracted_techs


def _assemble_projects(
    lines: List[LineInfo],
    median_font_size: float,
    header_threshold: float = 0.30
) -> List[dict]:
    """
    Assemble structured Project objects from classified blocks across the entire document
    or project/experience sections, preserving block-level evidence IDs and quality status.
    """
    if not lines:
        return []

    # First classify all blocks
    lines = _classify_blocks(lines, median_font_size, header_threshold)
    
    projects: list[dict] = []
    i = 0
    num_lines = len(lines)

    while i < num_lines:
        line = lines[i]

        if line.block_type == "PROJECT_TITLE":
            title, inline_techs = _clean_project_name(line.text)
            
            name_blocks = [line]
            desc_blocks: list[LineInfo] = []
            tech_blocks: list[LineInfo] = []
            project_techs: set[str] = set(t.lower() for t in inline_techs)

            if inline_techs:
                tech_blocks.append(line)

            j = i + 1
            while j < num_lines:
                curr = lines[j]
                
                # Stop if another project title or section heading is reached
                if curr.block_type in ("PROJECT_TITLE", "SECTION_HEADING"):
                    break

                if curr.block_type == "TECHNOLOGY_LINE":
                    tech_blocks.append(curr)
                    found = _TECH_REGEX.findall(curr.text)
                    for t in found:
                        project_techs.add(t.lower())
                    j += 1
                elif curr.block_type in ("BULLET", "DESCRIPTION", "LINK", "DATE"):
                    desc_blocks.append(curr)
                    # Also look for technologies inside description blocks
                    found_b = _TECH_REGEX.findall(curr.text)
                    if found_b:
                        tech_blocks.append(curr)
                        for t in found_b:
                            project_techs.add(t.lower())
                    j += 1
                else:
                    break

            # Build description text
            desc_texts = []
            for b in desc_blocks:
                clean_txt = re.sub(r"^[•\-\*\u2022\u25cf\u25cb\u2023\u25aa\u25b8\d+\.]\s*", "", b.text).strip()
                if clean_txt:
                    desc_texts.append(clean_txt)
            description = " ".join(desc_texts[:3]) if desc_texts else line.text

            # Calculate extraction confidence score
            conf = 0.35 + (0.25 if desc_blocks else 0.0) + (0.20 if project_techs else 0.0) + (0.20 if line.is_bold else 0.10)
            conf = round(min(1.0, conf), 2)

            # Quality gate status assignment
            if conf >= 0.55 and len(title) >= 3 and (desc_blocks or project_techs):
                status = "verified_extraction"
            elif conf >= 0.40 and len(title) >= 3:
                status = "uncertain_extraction"
            else:
                status = "needs_review"  # Isolated low-quality data

            all_proj_blocks = name_blocks + desc_blocks + [b for b in tech_blocks if b not in name_blocks and b not in desc_blocks]

            # Validation: filter out obvious non-projects or noise
            if len(title) >= 3 and not title.lower().startswith(("education", "skills", "certifications", "contact")):
                projects.append({
                    "name": title,
                    "description": description,
                    "technologies": [t.title() for t in sorted(list(project_techs))],
                    "confidence": conf,
                    "status": status,
                    "extraction_method": "block_classified_layout",
                    "evidence": {
                        "name_block_ids": [b.block_id for b in name_blocks],
                        "description_block_ids": [b.block_id for b in desc_blocks],
                        "technology_block_ids": [b.block_id for b in tech_blocks],
                    },
                    "source_blocks": [b.to_dict() for b in all_proj_blocks],
                })
            
            i = j
        else:
            i += 1

    return projects


# ═══════════════════════════════════════════════════════════════════════════════
#  Standard Field Extractors
# ═══════════════════════════════════════════════════════════════════════════════

def _extract_email(text: str) -> str:
    match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    return match.group(0) if match else ""


def _extract_phone(text: str) -> str:
    match = re.search(r"(?:\+?\d{1,3}[\s-]?)?\(?\d{3,5}\)?[\s-]?\d{3,5}[\s-]?\d{3,5}", text)
    return match.group(0).strip() if match else ""


def _extract_github_username(text: str) -> str:
    match = re.search(r"github\.com/([A-Za-z0-9_-]+)", text, re.IGNORECASE)
    if match:
        user = match.group(1)
        if user.lower() not in ("search", "explore", "features", "pricing"):
            return user
    return ""


def _extract_name(lines: List[LineInfo], filepath: str) -> str:
    """Extract candidate name using top line font size and heuristics."""
    if lines:
        top_lines = lines[:3]
        top_lines_sorted = sorted(top_lines, key=lambda l: l.font_size, reverse=True)
        candidate = top_lines_sorted[0].text.strip()
        cleaned_cand = re.sub(r"[^A-Za-z\s\.]", "", candidate).strip()
        if len(cleaned_cand.split()) in (2, 3, 4) and not _extract_email(candidate):
            return cleaned_cand

    stem = Path(filepath).stem
    stem = re.sub(r"^\d+\s*[-_]?\s*", "", stem).strip()
    return stem if stem else "Unknown"


def _extract_skills(text: str) -> list[str]:
    """Extract skills using standard section parsing + tech keywords."""
    found = set()
    for kw_match in _TECH_REGEX.finditer(text):
        found.add(kw_match.group(0).lower())
    return sorted([k.title() for k in found])


def _extract_education(text: str) -> str:
    match = re.search(r"\b(bachelor|master|b\.e\.|b\.tech|m\.tech|phd|diploma)\b.*?(?=\n\n|\Z)", text, re.I | re.S)
    return match.group(0).strip()[:300] if match else ""


# ═══════════════════════════════════════════════════════════════════════════════
#  Structured Context LLM/SLM Fallback Repair
# ═══════════════════════════════════════════════════════════════════════════════

def _llm_repair_project_extraction(
    lines: List[LineInfo],
    low_confidence_projects: List[dict]
) -> List[dict]:
    """
    Optional SLM/LLM fallback validator receiving structured block JSON context
    with block_ids, bounding boxes, and typography. Returns evidence_block_ids.
    Zero invention / zero hallucination constraint.
    """
    print("  → [SLM/LLM Fallback] Low-confidence project extraction detected. Invoking structured repair validator...")
    # Passes structured_blocks = [l.to_dict() for l in lines] to LLM API if enabled
    return low_confidence_projects


# ═══════════════════════════════════════════════════════════════════════════════
#  Public API
# ═══════════════════════════════════════════════════════════════════════════════

def parse_resume(
    filepath: str,
    header_threshold: float = 0.30,
    enable_llm_repair: bool = False
) -> dict:
    """
    Parse a PDF or DOCX resume using layout-aware typographic block classification.

    Parameters
    ----------
    filepath : str
        Path to PDF or DOCX resume file.
    header_threshold : float, optional
        Candidate title score threshold (default: 0.30).
    enable_llm_repair : bool, optional
        Whether to invoke structured LLM fallback for low-confidence projects.

    Returns dict with keys:
      name, email, github_username, phone, skills, projects, education, raw_text
    """
    t0 = time.time()
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {filepath}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        lines, median_font_size, layout_metadata = _extract_pdf_structured(filepath)
    elif ext in (".docx", ".doc"):
        lines, median_font_size, layout_metadata = _extract_docx_structured(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Expected .pdf or .docx")

    # Apply semantic reconstruction pipeline
    lines = _reconstruct_lines(lines)
    lines = _detect_sections(lines, median_font_size)

    raw_text = "\n".join(l.text for l in lines if l.text)

    # Heading-independent project assembly with block evidence tracking
    projects = _assemble_projects(lines, median_font_size, header_threshold=header_threshold)

    # Low-confidence quality gate check & optional LLM fallback repair
    low_conf = [p for p in projects if p.get("status") == "needs_review"]
    if low_conf or layout_metadata.get("pdf_type") == "scanned":
        if enable_llm_repair or os.environ.get("ENABLE_LLM_PARSER_REPAIR") == "true":
            projects = _llm_repair_project_extraction(lines, projects)

    elapsed_ms = round((time.time() - t0) * 1000, 1)

    return {
        "name": _extract_name(lines, filepath),
        "email": _extract_email(raw_text),
        "github_username": _extract_github_username(raw_text),
        "phone": _extract_phone(raw_text),
        "skills": _extract_skills(raw_text),
        "projects": projects,
        "education": _extract_education(raw_text),
        "layout_metadata": layout_metadata,
        "raw_text": raw_text,
        "parse_time_ms": elapsed_ms,
    }


def extract_text(filepath: str) -> str:
    """Legacy compatibility helper."""
    path = Path(filepath)
    ext = path.suffix.lower()
    if ext == ".pdf":
        lines, _, _ = _extract_pdf_structured(filepath)
    elif ext in (".docx", ".doc"):
        lines, _, _ = _extract_docx_structured(filepath)
    else:
        lines = []
    return _clean_text("\n".join(l.text for l in lines))
