"""
document_extractor.py - Lightweight physical text extraction for resumes.
Supports PDF (PyMuPDF) and DOCX (python-docx).
Provides both plain text extraction and numbered line extraction.
"""

import os
from typing import List, Tuple, Dict, Any
from collections import namedtuple

LineInfo = namedtuple("LineInfo", ["line_id", "text", "bbox", "page"])


def extract_plain_text(filepath: str) -> str:
    """
    Extracts raw, clean plain text from PDF or DOCX file without any line numbers.
    """
    lines, _ = extract_numbered_lines(filepath)
    return "\n".join([line.text for line in lines])


def extract_numbered_lines(filepath: str) -> Tuple[List[LineInfo], str]:
    """
    Extracts text from PDF/DOCX and returns a list of LineInfo objects
    along with a formatted string of text lines.
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(filepath)
    elif ext == ".docx":
        return _extract_docx(filepath)
    else:
        raise NotImplementedError(f"Unsupported file extension: {ext}")


def _extract_pdf(filepath: str) -> Tuple[List[LineInfo], str]:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF (fitz) is required for PDF extraction. pip install PyMuPDF")
        
    doc = fitz.open(filepath)
    all_lines = []
    line_id_counter = 1
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Sort blocks by vertical position to approximate reading order
        blocks = page.get_text("dict")["blocks"]
        blocks = [b for b in blocks if b["type"] == 0]
        blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
        
        for b in blocks:
            for l in b["lines"]:
                line_text = ""
                for s in l["spans"]:
                    text = s["text"].replace("\u2022", "-").strip()
                    if text:
                        line_text += text + " "
                        
                line_text = line_text.strip()
                if not line_text:
                    continue
                    
                bbox = l["bbox"]
                all_lines.append(LineInfo(
                    line_id=line_id_counter,
                    text=line_text,
                    bbox=bbox,
                    page=page_num + 1
                ))
                line_id_counter += 1
                
    doc.close()
    
    formatted_text = "\n".join([line.text for line in all_lines])
    return all_lines, formatted_text


def _extract_docx(filepath: str) -> Tuple[List[LineInfo], str]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX extraction. pip install python-docx")
        
    doc = Document(filepath)
    all_lines = []
    line_id_counter = 1
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        all_lines.append(LineInfo(
            line_id=line_id_counter,
            text=text,
            bbox=None,
            page=1
        ))
        line_id_counter += 1
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    all_lines.append(LineInfo(
                        line_id=line_id_counter,
                        text=text,
                        bbox=None,
                        page=1
                    ))
                    line_id_counter += 1
                    
    formatted_text = "\n".join([line.text for line in all_lines])
    return all_lines, formatted_text
