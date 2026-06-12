"""
HireFlow Resume Parser
-----------------------
Extracts text from PDF and DOCX resume files.
Uses PyMuPDF for PDFs and python-docx for Word documents.
"""

import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


def _clean_text(raw: str) -> str:
    """
    Normalize whitespace and strip non-printable characters.
    - Collapse runs of whitespace into single spaces.
    - Preserve paragraph boundaries (double newlines).
    - Strip leading / trailing whitespace.
    """
    # Replace tabs and form-feeds with spaces
    text = raw.replace("\t", " ").replace("\f", " ")
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse 3+ newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces (but not newlines) into one
    text = re.sub(r"[^\S\n]+", " ", text)
    # Strip each line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    # Final trim
    return text.strip()


def _extract_pdf(filepath: str) -> str:
    """Extract text from a PDF using PyMuPDF."""
    doc = fitz.open(filepath)
    pages: list[str] = []
    for page in doc:
        page_text = page.get_text("text")
        if page_text:
            pages.append(page_text)
    doc.close()
    return "\n\n".join(pages)


def _extract_docx(filepath: str) -> str:
    """Extract text from a DOCX using python-docx."""
    doc = Document(filepath)
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def extract_text(filepath: str) -> str:
    """
    Detect file type and extract text.

    Parameters
    ----------
    filepath : str
        Absolute or relative path to a PDF or DOCX file.

    Returns
    -------
    str
        Cleaned text content of the resume.

    Raises
    ------
    ValueError
        If the file extension is not .pdf or .docx.
    FileNotFoundError
        If the file does not exist.
    """
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {filepath}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        raw = _extract_pdf(filepath)
    elif ext in (".docx", ".doc"):
        raw = _extract_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Expected .pdf or .docx")

    cleaned = _clean_text(raw)
    if not cleaned:
        raise ValueError(f"No text could be extracted from: {filepath}")

    return cleaned
